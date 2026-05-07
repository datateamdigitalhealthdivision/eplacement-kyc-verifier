"""Build manuscript artifacts for the ePlacement KYC verifier v4.1 validation."""

from __future__ import annotations

import json
import math
import re
from pathlib import Path

import pandas as pd
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt
from PIL import Image, ImageDraw, ImageFont


HERE = Path(__file__).resolve().parent
WORKBOOK = Path(r"C:\Users\vivek\Downloads\eplacement_kyc_validation_analysis_v4_1.xlsx")
MANUSCRIPT_MD = HERE / "MANUSCRIPT.md"
OUTPUT_DOCX = HERE / "claim_guided_document_evidence_verifier_manuscript.docx"
TABLES_DIR = HERE / "tables"
FIGURES_DIR = HERE / "figures"
SUMMARY_JSON = HERE / "validation_summary.json"

PALETTE = {
    "ink": "#172033",
    "muted": "#5b667a",
    "line": "#c7d2e5",
    "panel": "#f5f8fc",
    "panel2": "#edf7f4",
    "purple": "#440154",
    "indigo": "#3B528B",
    "blue": "#31688E",
    "teal": "#21918C",
    "green": "#35B779",
    "lime": "#5DC863",
    "yellow": "#FDE725",
    "red": "#B63679",
    "white": "#ffffff",
}


def wilson_interval(successes: int, total: int, z: float = 1.96) -> dict[str, float]:
    if total <= 0:
        return {"estimate": 0.0, "lower": 0.0, "upper": 0.0}
    p = successes / total
    denom = 1 + z**2 / total
    center = (p + z**2 / (2 * total)) / denom
    half_width = z * math.sqrt((p * (1 - p) + z**2 / (4 * total)) / total) / denom
    return {"estimate": p, "lower": center - half_width, "upper": center + half_width}


def clean_key(value: object) -> str:
    return re.sub(r"[^a-z0-9]+", "_", str(value).strip().casefold()).strip("_")


def extract_validation_tables() -> None:
    """Extract small, de-identified validation summaries from the final workbook."""

    if not WORKBOOK.exists():
        raise FileNotFoundError(f"Validation workbook not found: {WORKBOOK}")

    TABLES_DIR.mkdir(parents=True, exist_ok=True)
    by_evidence = pd.read_excel(WORKBOOK, sheet_name="By evidence")
    comparison = pd.read_excel(WORKBOOK, sheet_name="v4 comparison")
    review_flag = pd.read_excel(WORKBOOK, sheet_name="Review flag")
    summary_raw = pd.read_excel(WORKBOOK, sheet_name="Summary", header=None)

    by_evidence.to_csv(TABLES_DIR / "by_evidence.csv", index=False)
    comparison.to_csv(TABLES_DIR / "v4_comparison.csv", index=False)
    review_flag.to_csv(TABLES_DIR / "review_flag.csv", index=False)

    summary_lookup: dict[str, object] = {}
    for _, row in summary_raw.iterrows():
        key = row.iloc[0]
        value = row.iloc[1] if len(row) > 1 else None
        if pd.notna(key) and pd.notna(value):
            summary_lookup[clean_key(key)] = value.item() if hasattr(value, "item") else value

    total_decisions = int(summary_lookup.get("binary_decisions", 0))
    tp = int(summary_lookup.get("tp", 0))
    fp = int(summary_lookup.get("fp", 0))
    tn = int(summary_lookup.get("tn", 0))
    fn = int(summary_lookup.get("fn", 0))
    applicants = int(summary_lookup.get("applicants_matched", 0))
    exact_matches = int(summary_lookup.get("exact_applicant_matches", 0))
    review_flagged = int(summary_lookup.get("review_flagged_applicants", 0))

    validation_summary = {
        "source_workbook": str(WORKBOOK),
        "model_evaluated": summary_lookup.get("model_evaluated"),
        "validation_set": summary_lookup.get("validation_set"),
        "applicants_matched": applicants,
        "binary_decisions": total_decisions,
        "confusion_matrix": {"tp": tp, "fp": fp, "tn": tn, "fn": fn},
        "metrics": {
            "accuracy": summary_lookup.get("accuracy"),
            "sensitivity_recall": summary_lookup.get("sensitivity_recall"),
            "specificity": summary_lookup.get("specificity"),
            "ppv_precision": summary_lookup.get("ppv_precision"),
            "npv": summary_lookup.get("npv"),
            "f1_score": summary_lookup.get("f1_score"),
            "exact_applicant_match_rate": summary_lookup.get("exact_applicant_match"),
            "review_rate": summary_lookup.get("review_rate"),
        },
        "descriptive_wilson_intervals": {
            "accuracy": wilson_interval(tp + tn, total_decisions),
            "sensitivity_recall": wilson_interval(tp, tp + fn),
            "specificity": wilson_interval(tn, tn + fp),
            "ppv_precision": wilson_interval(tp, tp + fp),
            "npv": wilson_interval(tn, tn + fn),
            "exact_applicant_match_rate": wilson_interval(exact_matches, applicants),
            "manual_review_rate": wilson_interval(review_flagged, applicants),
        },
        "conclusion": summary_lookup.get("conclusion"),
    }
    SUMMARY_JSON.write_text(json.dumps(validation_summary, indent=2), encoding="utf-8")


def font(size: int, bold: bool = False) -> ImageFont.FreeTypeFont | ImageFont.ImageFont:
    candidates = [
        r"C:\Windows\Fonts\aptos-bold.ttf" if bold else r"C:\Windows\Fonts\aptos.ttf",
        r"C:\Windows\Fonts\segoeuib.ttf" if bold else r"C:\Windows\Fonts\segoeui.ttf",
        r"C:\Windows\Fonts\arialbd.ttf" if bold else r"C:\Windows\Fonts\arial.ttf",
    ]
    for candidate in candidates:
        if Path(candidate).exists():
            return ImageFont.truetype(candidate, size=size)
    return ImageFont.load_default()


def wrap_text(draw: ImageDraw.ImageDraw, text: str, font_obj: ImageFont.ImageFont, max_width: int) -> list[str]:
    words = text.split()
    lines: list[str] = []
    current = ""
    for word in words:
        proposed = f"{current} {word}".strip()
        if draw.textbbox((0, 0), proposed, font=font_obj)[2] <= max_width or not current:
            current = proposed
        else:
            lines.append(current)
            current = word
    if current:
        lines.append(current)
    return lines


def draw_wrapped(
    draw: ImageDraw.ImageDraw,
    text: str,
    xy: tuple[int, int],
    max_width: int,
    font_obj: ImageFont.ImageFont,
    fill: str = PALETTE["ink"],
    line_gap: int = 8,
) -> int:
    x, y = xy
    for line in wrap_text(draw, text, font_obj, max_width):
        draw.text((x, y), line, font=font_obj, fill=fill)
        y += font_obj.size + line_gap
    return y


def arrow(draw: ImageDraw.ImageDraw, start: tuple[int, int], end: tuple[int, int], fill: str = PALETTE["line"], width: int = 5) -> None:
    draw.line([start, end], fill=fill, width=width)
    x1, y1 = start
    x2, y2 = end
    if x2 >= x1:
        head = [(x2, y2), (x2 - 18, y2 - 12), (x2 - 18, y2 + 12)]
    else:
        head = [(x2, y2), (x2 + 18, y2 - 12), (x2 + 18, y2 + 12)]
    draw.polygon(head, fill=fill)


def card(
    draw: ImageDraw.ImageDraw,
    xyxy: tuple[int, int, int, int],
    title: str,
    body: str,
    accent: str,
    fill: str = PALETTE["white"],
    title_size: int = 34,
    body_size: int = 26,
) -> None:
    x1, y1, x2, y2 = xyxy
    draw.rounded_rectangle(xyxy, radius=32, fill=fill, outline=PALETTE["line"], width=3)
    draw.rounded_rectangle((x1, y1, x1 + 18, y2), radius=12, fill=accent)
    draw.text((x1 + 45, y1 + 34), title, font=font(title_size, True), fill=PALETTE["ink"])
    draw_wrapped(draw, body, (x1 + 45, y1 + 88), x2 - x1 - 90, font(body_size), fill=PALETTE["muted"], line_gap=7)


def canvas(title: str, subtitle: str | None = None, size: tuple[int, int] = (2200, 1250)) -> tuple[Image.Image, ImageDraw.ImageDraw]:
    image = Image.new("RGB", size, PALETTE["white"])
    draw = ImageDraw.Draw(image)
    draw.rectangle((0, 0, size[0], 140), fill="#ecf5ff")
    draw.text((80, 42), title, font=font(44, True), fill=PALETTE["ink"])
    if subtitle:
        draw.text((82, 96), subtitle, font=font(24), fill=PALETTE["muted"])
    return image, draw


def save_figure(image: Image.Image, name: str) -> None:
    FIGURES_DIR.mkdir(parents=True, exist_ok=True)
    image.save(FIGURES_DIR / name, optimize=True)


def create_figures() -> None:
    by_evidence = pd.read_excel(WORKBOOK, sheet_name="By evidence")
    summary = json.loads(SUMMARY_JSON.read_text(encoding="utf-8"))

    # Figure 1: problem and contribution.
    image, draw = canvas(
        "Administrative evidence verification problem",
        "The method supports people applying for human-resource placement where claims must be backed by uploaded proof.",
    )
    boxes = [
        ("Applicant claims", "Marital, health, family, spouse-location, disability and examination claims are entered as structured form data.", PALETTE["blue"]),
        ("Document bundles", "Applicants upload heterogeneous PDFs: scanned certificates, letters, medical notes, forms and multi-page mixed evidence.", PALETTE["lime"]),
        ("Verifier method", "A local-first harness checks only claimed categories, combining rules, OCR, page images and a vision-language model.", PALETTE["teal"]),
        ("Operator queue", "Outputs are tick sheets, proof summaries, page references and a targeted manual-review queue.", PALETTE["green"]),
    ]
    x = 80
    for i, (title, body, colour) in enumerate(boxes):
        card(draw, (x, 230, x + 480, 580), title, body, colour, fill=PALETTE["panel"])
        if i < len(boxes) - 1:
            arrow(draw, (x + 500, 405), (x + 590, 405), fill="#9db1d1")
        x += 530
    metrics = [
        ("238", "applicants"),
        ("1,428", "evidence decisions"),
        ("6", "evidence domains"),
        ("26.9%", "manual-review queue"),
    ]
    for i, (value, label) in enumerate(metrics):
        cx = 280 + i * 520
        draw.ellipse((cx - 115, 735, cx + 115, 965), fill="#eef8f6", outline="#9bd8cd", width=4)
        w = draw.textbbox((0, 0), value, font=font(46, True))[2]
        draw.text((cx - w / 2, 805), value, font=font(46, True), fill=PALETTE["teal"])
        lw = draw.textbbox((0, 0), label, font=font(23))[2]
        draw.text((cx - lw / 2, 870), label, font=font(23), fill=PALETTE["muted"])
    draw_wrapped(
        draw,
        "Contribution: a claim-guided proof verifier that treats accuracy, auditability, privacy and review workload as one coupled implementation problem.",
        (170, 1050),
        1850,
        font(30, True),
        fill=PALETTE["ink"],
    )
    save_figure(image, "figure_1_problem_contribution.png")

    # Figure 2: evidence taxonomy.
    image, draw = canvas("Six proof categories used by the verifier", "Evidence is only checked when the applicant has claimed that category.")
    taxonomy = [
        ("Marriage", "Marriage or nikah certificate; clear spouse relationship proof.", PALETTE["blue"]),
        ("Self illness", "Medical proof where the patient is the applicant or strongly implied to be the applicant.", PALETTE["teal"]),
        ("Family illness", "Medical proof for spouse, child, parent, dependent or other relevant family member.", PALETTE["lime"]),
        ("Spouse location", "Spouse workplace, posting, residence or location evidence relevant to placement.", PALETTE["purple"]),
        ("OKU self/family", "Official OKU card, JKM document, disability registration or clear disability evidence.", PALETTE["green"]),
        ("MedEX / exam", "MedEX, GCFM, postgraduate or specialist exam registration, result, attendance or certificate; not routine physical examination.", PALETTE["yellow"]),
    ]
    for i, (title, body, colour) in enumerate(taxonomy):
        row = i // 3
        col = i % 3
        x1 = 90 + col * 700
        y1 = 230 + row * 390
        card(draw, (x1, y1, x1 + 620, y1 + 300), title, body, colour, fill="#fbfcff", title_size=32, body_size=25)
    draw.rounded_rectangle((250, 1025, 1950, 1140), radius=30, fill="#fff7ed", outline="#fed7aa", width=3)
    draw_wrapped(
        draw,
        "Design rule: unclaimed categories are not labelled as positive evidence. This reduces false positives from broad document classification.",
        (305, 1060),
        1580,
        font(28, True),
        fill="#9a3412",
    )
    save_figure(image, "figure_2_evidence_taxonomy.png")

    # Figure 3: workflow.
    image, draw = canvas("Claim-guided document evidence workflow", "Operational pipeline from applicant form to auditable decision outputs.", size=(2300, 1350))
    steps = [
        ("1. Form ingest", "Column mapping and row normalisation"),
        ("2. Claim extraction", "Convert structured fields into claimed_* booleans"),
        ("3. PDF acquisition", "Use original link or cached local document"),
        ("4. Render and OCR", "Direct text first; OCR for scanned pages"),
        ("5. Page signals", "Keywords, names, ICs, layout cues and page images"),
        ("6. AI harness", "Local VLM verifies claimed proof only"),
        ("7. Exports", "Tick sheet, scoring sheet, review queue and audit trail"),
    ]
    x = 70
    y = 270
    for i, (title, body) in enumerate(steps):
        card(draw, (x, y, x + 285, y + 310), title, body, [PALETTE["blue"], PALETTE["teal"], PALETTE["lime"]][i % 3], fill=PALETTE["panel"], title_size=27, body_size=21)
        if i < len(steps) - 1:
            arrow(draw, (x + 300, y + 155), (x + 355, y + 155), fill="#a8b8d6", width=4)
        x += 315
    draw.rounded_rectangle((180, 760, 2120, 1120), radius=36, fill="#f8fafc", outline="#cbd5e1", width=3)
    draw.text((240, 815), "Safeguards embedded across the workflow", font=font(34, True), fill=PALETTE["ink"])
    safeguards = [
        "Skip unclaimed categories",
        "Cache OCR and model outputs",
        "Second-pass targeted verification",
        "Structured JSON parsing",
        "Proof-strength scoring",
        "Original PDF links retained",
        "Manual review triggers",
    ]
    for i, item in enumerate(safeguards):
        x0 = 255 + (i % 4) * 465
        y0 = 900 + (i // 4) * 90
        draw.rounded_rectangle((x0, y0, x0 + 395, y0 + 58), radius=20, fill="#e0f2fe", outline="#bae6fd")
        draw.text((x0 + 24, y0 + 16), item, font=font(21, True), fill="#075985")
    save_figure(image, "figure_3_workflow.png")

    # Figure 4: harness architecture.
    image, draw = canvas("AI harness and local inference architecture", "The cool part: model calls are constrained by deterministic context, schema and audit policy.", size=(2300, 1400))
    layers = [
        ("Langflow orchestration layer", "Readable workflow graph coordinates loader, fetcher, OCR router, verifier and export writer.", PALETTE["blue"]),
        ("Deterministic guardrails", "Claim extraction, unclaimed-category skipping, keyword priors, proof rules and confidence thresholds.", PALETTE["teal"]),
        ("Document perception layer", "PDF rendering, direct text extraction, Tesseract route, PaddleOCR fallback and page-image retention.", PALETTE["lime"]),
        ("Local multimodal model layer", "Ollama serves Qwen2.5-VL locally for page-level visual verification of claimed evidence.", PALETTE["purple"]),
        ("Structured adjudication layer", "JSON schema parsing, proof_strength, supporting_page, evidence_summary and check_required policy.", PALETTE["green"]),
    ]
    for i, (title, body, colour) in enumerate(layers):
        y1 = 220 + i * 205
        card(draw, (140, y1, 1580, y1 + 145), title, body, colour, fill="#fbfdff", title_size=31, body_size=24)
    for i in range(4):
        arrow(draw, (860, 370 + i * 205), (860, 425 + i * 205), fill="#94a3b8", width=5)
    draw.rounded_rectangle((1680, 270, 2150, 1050), radius=40, fill="#f0fdfa", outline="#99f6e4", width=4)
    draw.text((1735, 330), "Why this matters", font=font(34, True), fill=PALETTE["teal"])
    bullets = [
        "No external document API required",
        "Prompt sees applicant claims",
        "Model cannot reward unclaimed categories",
        "Outputs are auditable, not just labels",
        "Operator can inspect source PDF immediately",
    ]
    y = 410
    for item in bullets:
        draw.ellipse((1735, y + 9, 1755, y + 29), fill=PALETTE["teal"])
        y = draw_wrapped(draw, item, (1775, y), 310, font(25), fill=PALETTE["ink"], line_gap=7) + 22
    save_figure(image, "figure_4_ai_harness.png")

    # Figure 5: broad vs claim-guided.
    image, draw = canvas("From broad classification to claim-guided proof verification", "The methodological shift is from guessing document type to testing applicant-declared claims.", size=(2200, 1250))
    draw.rounded_rectangle((120, 245, 1010, 1000), radius=42, fill="#fff1f2", outline="#fecdd3", width=4)
    draw.rounded_rectangle((1190, 245, 2080, 1000), radius=42, fill="#ecfdf5", outline="#bbf7d0", width=4)
    draw.text((180, 310), "Broad classifier", font=font(40, True), fill=PALETTE["red"])
    draw.text((1250, 310), "Claim-guided verifier", font=font(40, True), fill=PALETTE["green"])
    broad = [
        "Scans the whole PDF for all six classes",
        "May label plausible but unclaimed evidence",
        "More false positives from generic letters or medical pages",
        "Useful for discovery, weaker for proof checking",
    ]
    guided = [
        "Reads applicant form first",
        "Only verifies categories the applicant claimed",
        "Stops when claimed proof is sufficiently supported",
        "Designed for triage and audit, not autonomous rejection",
    ]
    y = 410
    for item in broad:
        draw.text((185, y), "x", font=font(30, True), fill=PALETTE["red"])
        y = draw_wrapped(draw, item, (230, y), 690, font(28), fill=PALETTE["ink"]) + 32
    y = 410
    for item in guided:
        draw.text((1255, y), "+", font=font(30, True), fill=PALETTE["green"])
        y = draw_wrapped(draw, item, (1300, y), 690, font(28), fill=PALETTE["ink"]) + 32
    draw.rounded_rectangle((760, 555, 1440, 695), radius=34, fill="#eff6ff", outline="#bfdbfe", width=3)
    draw.text((845, 600), "Methodological contribution", font=font(34, True), fill=PALETTE["blue"])
    save_figure(image, "figure_5_method_shift.png")

    # Figure 6: evidence type performance as a scientific grouped bar chart.
    image = Image.new("RGB", (2300, 1300), PALETTE["white"])
    draw = ImageDraw.Draw(image)
    metrics = [("Sensitivity", "Sensitivity", PALETTE["purple"]), ("Specificity", "Specificity", PALETTE["teal"]), ("F1", "F1 score", PALETTE["yellow"])]
    label_map = {
        "marriage": "Marriage",
        "self_illness": "Self illness",
        "family_illness": "Family illness",
        "spouse_location": "Spouse location",
        "oku_self_or_family": "OKU self/family",
        "medex_other_exam": "MedEX/other exam",
    }
    labels = [label_map.get(str(x), str(x).replace("_", " ").title()) for x in by_evidence["Evidence type"]]
    chart_left, chart_top = 560, 130
    chart_width, chart_height = 1450, 860
    chart_bottom = chart_top + chart_height
    for tick in range(0, 101, 20):
        x = chart_left + int(chart_width * tick / 100)
        draw.line((x, chart_top, x, chart_bottom), fill="#e2e8f0", width=2)
        draw.text((x - 24, chart_bottom + 25), f"{tick}", font=font(22), fill=PALETTE["muted"])
    draw.line((chart_left, chart_bottom, chart_left + chart_width, chart_bottom), fill=PALETTE["ink"], width=3)
    draw.line((chart_left, chart_top, chart_left, chart_bottom), fill=PALETTE["ink"], width=3)
    row_gap = chart_height / len(labels)
    bar_h = 24
    for idx, label in enumerate(labels):
        row_mid = chart_top + int(row_gap * (idx + 0.5))
        draw.text((90, row_mid - 18), label, font=font(27, True), fill=PALETTE["ink"])
        for m_idx, (col, _, colour) in enumerate(metrics):
            value = float(by_evidence.iloc[idx][col]) * 100
            y = row_mid - 42 + m_idx * 30
            x2 = chart_left + int(chart_width * value / 100)
            draw.rectangle((chart_left, y, x2, y + bar_h), fill=colour)
            draw.text((x2 + 12, y - 3), f"{value:.1f}", font=font(20), fill=PALETTE["ink"])
    draw.text((chart_left + chart_width // 2 - 110, chart_bottom + 75), "Performance (%)", font=font(25, True), fill=PALETTE["ink"])
    legend_x = 1180
    for idx, (_, label, colour) in enumerate(metrics):
        x = legend_x + idx * 270
        draw.rectangle((x, 48, x + 42, 72), fill=colour)
        draw.text((x + 55, 42), label, font=font(24), fill=PALETTE["ink"])
    save_figure(image, "figure_6_performance_by_evidence.png")

    # Figure 7: manual review time as a scientific workload chart.
    image = Image.new("RGB", (2200, 1200), PALETTE["white"])
    draw = ImageDraw.Draw(image)
    applicants = int(summary["applicants_matched"])
    review = int(round(float(summary["metrics"]["review_rate"]) * applicants))
    manual_rate_per_hour = 5
    full_manual_hours = applicants / manual_rate_per_hour
    assisted_hours = review / manual_rate_per_hour
    avoided_hours = full_manual_hours - assisted_hours
    values = [
        ("Full manual review", full_manual_hours, PALETTE["purple"]),
        ("AI-assisted review queue", assisted_hours, PALETTE["teal"]),
        ("Reviewer-hours avoided", avoided_hours, PALETTE["yellow"]),
    ]
    chart_left, chart_top = 280, 90
    chart_width, chart_height = 1580, 760
    chart_bottom = chart_top + chart_height
    max_value = 55
    for tick in range(0, 56, 10):
        y = chart_bottom - int(chart_height * tick / max_value)
        draw.line((chart_left, y, chart_left + chart_width, y), fill="#e2e8f0", width=2)
        draw.text((chart_left - 70, y - 14), f"{tick}", font=font(22), fill=PALETTE["muted"])
    draw.line((chart_left, chart_bottom, chart_left + chart_width, chart_bottom), fill=PALETTE["ink"], width=3)
    draw.line((chart_left, chart_top, chart_left, chart_bottom), fill=PALETTE["ink"], width=3)
    bar_w = 280
    spacing = 260
    for idx, (label, value, colour) in enumerate(values):
        x1 = chart_left + 210 + idx * (bar_w + spacing)
        y1 = chart_bottom - int(chart_height * value / max_value)
        draw.rectangle((x1, y1, x1 + bar_w, chart_bottom), fill=colour)
        draw.text((x1 + 54, y1 - 48), f"{value:.1f}", font=font(31, True), fill=PALETTE["ink"])
        draw_wrapped(draw, label, (x1 - 40, chart_bottom + 38), bar_w + 80, font(23, True), fill=PALETTE["ink"], line_gap=6)
    draw.text((chart_left, chart_top - 55), "Reviewer time (person-hours)", font=font(25, True), fill=PALETTE["ink"])
    save_figure(image, "figure_7_review_funnel.png")


def split_table_row(line: str) -> list[str]:
    return [cell.strip() for cell in line.strip().strip("|").split("|")]


def is_table_separator(line: str) -> bool:
    cells = split_table_row(line)
    return bool(cells) and all(re.fullmatch(r":?-{3,}:?", cell) for cell in cells)


def add_markdown_table(document: Document, lines: list[str]) -> None:
    rows = [split_table_row(line) for line in lines if not is_table_separator(line)]
    if not rows:
        return
    width = max(len(row) for row in rows)
    table = document.add_table(rows=len(rows), cols=width)
    table.style = "Table Grid"
    for row_idx, row in enumerate(rows):
        for col_idx in range(width):
            cell_text = row[col_idx] if col_idx < len(row) else ""
            paragraph = table.cell(row_idx, col_idx).paragraphs[0]
            run = paragraph.add_run(cell_text)
            if row_idx == 0:
                run.bold = True
            if col_idx > 0:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    document.add_paragraph()


def add_paragraph_with_basic_formatting(document: Document, text: str, style: str | None = None) -> None:
    paragraph = document.add_paragraph(style=style)
    cursor = 0
    for match in re.finditer(r"\*\*(.+?)\*\*", text):
        if match.start() > cursor:
            paragraph.add_run(text[cursor : match.start()])
        bold_run = paragraph.add_run(match.group(1))
        bold_run.bold = True
        cursor = match.end()
    if cursor < len(text):
        paragraph.add_run(text[cursor:])


def build_docx() -> None:
    document = Document()
    section = document.sections[0]
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)

    styles = document.styles
    styles["Normal"].font.name = "Aptos"
    styles["Normal"].font.size = Pt(10)
    styles["Normal"].paragraph_format.space_after = Pt(6)
    styles["Normal"].paragraph_format.line_spacing = 1.08
    for style_name, size in [("Heading 1", 16), ("Heading 2", 13), ("Heading 3", 11)]:
        styles[style_name].font.name = "Aptos"
        styles[style_name].font.size = Pt(size)

    lines = MANUSCRIPT_MD.read_text(encoding="utf-8").splitlines()
    table_buffer: list[str] = []

    def flush_table() -> None:
        nonlocal table_buffer
        if table_buffer:
            add_markdown_table(document, table_buffer)
            table_buffer = []

    for raw_line in lines:
        line = raw_line.rstrip()
        if line.startswith("|"):
            table_buffer.append(line)
            continue
        flush_table()
        if not line.strip():
            continue
        if re.match(r"^Table\s+\d", line):
            caption_para = document.add_paragraph()
            caption_para.paragraph_format.space_before = Pt(8)
            caption_para.paragraph_format.space_after = Pt(3)
            caption_run = caption_para.add_run(line)
            caption_run.bold = True
            caption_run.italic = True
            caption_run.font.size = Pt(9)
            continue
        image_match = re.match(r"^!\[(.*?)\]\((.*?)\)$", line)
        if image_match:
            caption = image_match.group(1).strip()
            image_path = (HERE / image_match.group(2).strip()).resolve()
            if image_path.exists():
                paragraph = document.add_paragraph()
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = paragraph.add_run()
                run.add_picture(str(image_path), width=Inches(6.8))
                if caption:
                    caption_para = document.add_paragraph()
                    caption_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    caption_run = caption_para.add_run(caption)
                    caption_run.italic = True
                    caption_run.font.size = Pt(9)
            continue
        if line.startswith("# "):
            document.add_heading(line[2:].strip(), level=0)
        elif line.startswith("## "):
            document.add_heading(line[3:].strip(), level=1)
        elif line.startswith("### "):
            document.add_heading(line[4:].strip(), level=2)
        elif line.startswith("- "):
            add_paragraph_with_basic_formatting(document, line[2:].strip(), style="List Bullet")
        elif re.match(r"^\d+\.\s+", line):
            add_paragraph_with_basic_formatting(document, re.sub(r"^\d+\.\s+", "", line), style="List Number")
        else:
            add_paragraph_with_basic_formatting(document, line)
    flush_table()

    document.save(OUTPUT_DOCX)


def main() -> None:
    extract_validation_tables()
    create_figures()
    build_docx()
    print(f"Wrote {OUTPUT_DOCX}")
    print(f"Wrote {SUMMARY_JSON}")
    print(f"Wrote CSV tables to {TABLES_DIR}")


if __name__ == "__main__":
    main()
